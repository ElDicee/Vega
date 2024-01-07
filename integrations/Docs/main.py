import os.path

from integrations import VegaAPI as api
from docx import Document


class DocInstance():
    def __init__(self):
        self.document = Document()


def create_instance():
    return DocInstance()


def add_paragraph(instance, text: str):
    instance.document.add_paragraph(text)
    return instance


def save_document(instance, folder_path: str, filename: str):
    instance.document.save(os.path.join(folder_path, f"{filename}.docx"))
    return instance


def add_picture(instance, image_path: str):
    instance.document.add_picture(image_path)
    return instance


def vega_main():
    vega = api.Vega_Portal()
    vega.set_name("Docs ITG")
    vega.add_method(
        api.Method(create_instance, api.EXECUTION, formal_name="Create Instance", outputs={"Instance": None},
                   event=False))
    vega.add_method(
        api.Method(add_paragraph, api.EXECUTION, formal_name="Add Paragraph", outputs={"Instance": None}, event=False))
    vega.add_method(api.Method(save_document, api.EXECUTION, formal_name="Save Document", event=False))
    return vega
